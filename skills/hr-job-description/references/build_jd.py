"""
Locked template for PIF-styled Job Descriptions.

Called by the hr-job-description skill (v1.9.0+) to guarantee that every JD
produced from this skill has pixel-identical formatting — same fonts, sizes,
colors, indents, table widths, bullet alignment, and section spacing.

Do NOT hand-generate docx code inside SKILL.md. Instead:
  1. Assemble the JD content into a JSON blob matching the schema below.
  2. Invoke this script with the JSON path as its first argument, OR pipe
     the JSON to stdin.
  3. The script writes the final .docx to
     ~/HR-Workspace/hr-job-description/outputs/YYYYMMDD_JD_<Role_Slug>.docx
     with automatic file-lock retry (_v2, _v3, ..._v9) if the target is
     open in Word.

JSON schema (all fields required unless noted):

{
  "role": {
    "name": "Senior Analyst",               // used in title + filename
    "level": "Senior",                       // shown in subtitle
    "portfolio_or_division": "AI & Digital Hardware",
                                             // verbatim glossary term
    "employment_type": "Full-time (permanent)",
    "work_arrangement": "On-site (Riyadh)",
    "location": "Riyadh",
    "reports_to": "Head of AI & Digital Hardware"
  },
  "role_overview": ["paragraph 1", "paragraph 2"],   // 1-2 short paragraphs
  "responsibilities": ["bullet 1", ...],              // 5-8 bullets
  "must_haves": ["bullet 1", ...],                    // 4-6 bullets
  "nice_to_haves": ["bullet 1", ...],                 // 3-5 bullets
  "day_30": ["bullet 1", "bullet 2"],
  "day_60": ["bullet 1", "bullet 2"],
  "day_90": ["bullet 1", "bullet 2"],
  "what_we_offer": "paragraph"
}
"""
from __future__ import annotations

import json
import os
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ---------- Locked style tokens ----------

PIF_GREEN = "005C4D"
TAN = "C4984F"
TEXT_GRAY = "595959"
SOFT_GRAY = "9A9A9A"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"

FONT_PRIMARY = "Fund Light"
FONT_FALLBACK = "Calibri"

# Locked sizes and spacing
TITLE_PT = 20
SUBTITLE_PT = 12
SECTION_HEADING_PT = 14
BODY_PT = 11
BULLET_PT = 11
TABLE_HEADER_PT = 11
CALLOUT_BULLET_PT = 11
FOOTER_PT = 8

# Locked table widths (Cm)
WORK_DETAILS_COL_WIDTHS = [Cm(4.2), Cm(4.2), Cm(4.2), Cm(4.2)]  # 4 equal
DAY_30_60_90_COL_WIDTHS = [Cm(5.6), Cm(5.6), Cm(5.6)]           # 3 equal
CALLOUT_TABLE_WIDTH = Cm(16.8)                                   # single column
CALLOUT_BULLET_INDENT = Cm(0.4)


# ---------- Low-level helpers (identical to retention skill) ----------

def _set_run_font(run, name: str = FONT_PRIMARY, size: int = BODY_PT,
                  color_hex: str = TEXT_GRAY, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    run.bold = bold
    run.italic = italic


def _add_para(doc, text: str, size: int = BODY_PT, color: str = TEXT_GRAY,
              bold: bool = False, italic: bool = False, align=None, space_after: int = 6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    _set_run_font(run, size=size, color_hex=color, bold=bold, italic=italic)
    return p


def _add_heading(doc, text: str, size: int = SECTION_HEADING_PT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, size=size, color_hex=PIF_GREEN, bold=True)
    return p


def _add_bullet(doc, text: str, size: int = BULLET_PT, color: str = TEXT_GRAY):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    _set_run_font(run, size=size, color_hex=color)
    return p


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = SOFT_GRAY, size_eighths: int = 4) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size_eighths))
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _horizontal_rule(doc, color: str = PIF_GREEN) -> None:
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


def _save_with_retry(doc, base_path: Path) -> Path:
    """Save doc, auto-retrying with _v2, _v3, ... if the target file is open."""
    candidates = [base_path] + [
        base_path.with_name(f"{base_path.stem}_v{i}{base_path.suffix}")
        for i in range(2, 10)
    ]
    for p in candidates:
        try:
            doc.save(str(p))
            return p
        except PermissionError:
            continue
    raise PermissionError(f"All {len(candidates)} candidate paths for {base_path.name} are locked.")


# ---------- Section builders (each is deterministic) ----------

def _build_header(doc, role: dict) -> None:
    """Title + subtitle + horizontal rule."""
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run_t = title.add_run(f"Job Description — {role['name']}")
    _set_run_font(run_t, size=TITLE_PT, color_hex=PIF_GREEN, bold=True)

    subtitle_text = f"{role['portfolio_or_division']} · {role['level']} · Alat"
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(4)
    run_s = subtitle.add_run(subtitle_text)
    _set_run_font(run_s, size=SUBTITLE_PT, color_hex=TEXT_GRAY, italic=True)

    _horizontal_rule(doc)


def _build_work_details(doc, role: dict) -> None:
    """4-column work details block (Employment type · Work arrangement · Location · Reports to)."""
    headers = ["Employment type", "Work arrangement", "Location", "Reports to"]
    values = [
        role["employment_type"],
        role["work_arrangement"],
        role["location"],
        role["reports_to"],
    ]
    tbl = doc.add_table(rows=2, cols=4)
    for col_idx, w in enumerate(WORK_DETAILS_COL_WIDTHS):
        for row in tbl.rows:
            row.cells[col_idx].width = w

    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        _shade_cell(cell, PIF_GREEN)
        _set_cell_borders(cell, color=PIF_GREEN, size_eighths=6)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(h)
        _set_run_font(run, size=TABLE_HEADER_PT, color_hex=WHITE, bold=True)

    for i, v in enumerate(values):
        cell = tbl.rows[1].cells[i]
        _shade_cell(cell, LIGHT_GRAY)
        _set_cell_borders(cell, color=SOFT_GRAY, size_eighths=4)
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(v)
        _set_run_font(run, size=10, color_hex=TEXT_GRAY)

    doc.add_paragraph()  # spacer


def _build_role_overview(doc, paragraphs: list) -> None:
    _add_heading(doc, "Role Overview")
    for para in paragraphs:
        _add_para(doc, para, space_after=8)


def _build_responsibilities(doc, bullets: list) -> None:
    _add_heading(doc, "Key Responsibilities")
    for b in bullets:
        _add_bullet(doc, b)


def _build_must_haves(doc, bullets: list) -> None:
    _add_heading(doc, "Must-Have Qualifications")
    for b in bullets:
        _add_bullet(doc, b)


def _build_nice_to_haves(doc, bullets: list) -> None:
    """Tan-bordered callout box — single 1x1 table, all bullets inside the cell with identical left indent.

    Locked spec (from SKILL.md §Styling):
      - Single 1x1 table with tan (C4984F) 1.5pt border on all four sides
      - First paragraph inside the cell is a bullet (no leading blank paragraph)
      - Every bullet paragraph has List Bullet style AND explicit left_indent = CALLOUT_BULLET_INDENT
    """
    _add_heading(doc, "Nice-to-Have Qualifications")

    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    cell.width = CALLOUT_TABLE_WIDTH

    # Tan border, ~1.5pt (size 12 = 1.5pt in Word's eighths-of-a-point units)
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "12")
        b.set(qn("w:color"), TAN)
        tcBorders.append(b)
    tcPr.append(tcBorders)

    # Cell padding
    tcMar = OxmlElement("w:tcMar")
    for edge, val in [("top", "120"), ("left", "160"), ("bottom", "120"), ("right", "160")]:
        m = OxmlElement(f"w:{edge}")
        m.set(qn("w:w"), val)
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)

    # Replace the cell's default paragraph with the first bullet
    cell.text = ""
    first_p = cell.paragraphs[0]
    first_p.style = doc.styles["List Bullet"]
    first_p.paragraph_format.left_indent = CALLOUT_BULLET_INDENT
    first_p.paragraph_format.space_after = Pt(3)

    if not bullets:
        # No nice-to-haves supplied — emit a single italic placeholder so the callout isn't empty
        run = first_p.add_run("(none)")
        _set_run_font(run, size=CALLOUT_BULLET_PT, color_hex=TEXT_GRAY, italic=True)
        return

    # First bullet reuses the cell's default paragraph (this is the fix for the v1.6.0 misalignment bug)
    run = first_p.add_run(bullets[0])
    _set_run_font(run, size=CALLOUT_BULLET_PT, color_hex=TEXT_GRAY)

    # Subsequent bullets are new paragraphs INSIDE the cell, same style and indent
    for b in bullets[1:]:
        p = cell.add_paragraph()
        p.style = doc.styles["List Bullet"]
        p.paragraph_format.left_indent = CALLOUT_BULLET_INDENT
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(b)
        _set_run_font(run, size=CALLOUT_BULLET_PT, color_hex=TEXT_GRAY)


def _build_30_60_90(doc, day_30: list, day_60: list, day_90: list) -> None:
    """3-column table for 30/60/90-day success expectations."""
    _add_heading(doc, "30 / 60 / 90-Day Success Expectations")

    tbl = doc.add_table(rows=2, cols=3)
    for col_idx, w in enumerate(DAY_30_60_90_COL_WIDTHS):
        for row in tbl.rows:
            row.cells[col_idx].width = w

    # Header row
    for i, h in enumerate(["First 30 Days", "First 60 Days", "First 90 Days"]):
        cell = tbl.rows[0].cells[i]
        _shade_cell(cell, PIF_GREEN)
        _set_cell_borders(cell, color=PIF_GREEN, size_eighths=6)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_run_font(run, size=TABLE_HEADER_PT, color_hex=WHITE, bold=True)

    # Body row — bullets stacked vertically inside each cell
    for col_i, bullets in enumerate([day_30, day_60, day_90]):
        cell = tbl.rows[1].cells[col_i]
        fill = WHITE if col_i % 2 == 0 else LIGHT_GRAY
        _shade_cell(cell, fill)
        _set_cell_borders(cell, color=SOFT_GRAY, size_eighths=4)
        cell.text = ""
        for j, b in enumerate(bullets):
            if j == 0:
                p = cell.paragraphs[0]
            else:
                p = cell.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Cm(0.25)
            run_bullet = p.add_run("• ")
            _set_run_font(run_bullet, size=10, color_hex=PIF_GREEN, bold=True)
            run_text = p.add_run(b)
            _set_run_font(run_text, size=10, color_hex=TEXT_GRAY)


def _build_what_we_offer(doc, text: str) -> None:
    _add_heading(doc, "What We Offer")
    _add_para(doc, text, space_after=8)


def _build_footer(doc) -> None:
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.text = ""
    run = fp.add_run("Alat Talent Acquisition")
    _set_run_font(run, size=FOOTER_PT, color_hex=SOFT_GRAY, italic=True)


# ---------- Public entry point ----------

def build_jd(data: dict, out_path: Path) -> Path:
    """Build the JD document from a structured JSON blob. Deterministic — every run
    with the same input produces byte-identical output modulo the docx timestamp."""
    doc = Document()

    # Locked page margins
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    _build_header(doc, data["role"])
    _build_work_details(doc, data["role"])
    _build_role_overview(doc, data["role_overview"])
    _build_responsibilities(doc, data["responsibilities"])
    _build_must_haves(doc, data["must_haves"])
    _build_nice_to_haves(doc, data.get("nice_to_haves") or [])
    _build_30_60_90(doc, data["day_30"], data["day_60"], data["day_90"])
    _build_what_we_offer(doc, data["what_we_offer"])
    _build_footer(doc)

    return _save_with_retry(doc, out_path)


def _slug(name: str) -> str:
    return "".join(ch if ch.isalnum() else "_" for ch in name).strip("_")


def _default_output_path(role_name: str) -> Path:
    today = datetime.now().strftime("%Y%m%d")
    out_dir = Path(os.path.expanduser("~")) / "HR-Workspace" / "hr-job-description" / "outputs"
    out_dir.mkdir(parents=True, exist_ok=True)
    return out_dir / f"{today}_JD_{_slug(role_name)}.docx"


def _load_input() -> dict:
    # utf-8-sig strips a BOM if present (Windows PowerShell writes utf-8 with BOM by default).
    if len(sys.argv) >= 2:
        with open(sys.argv[1], "r", encoding="utf-8-sig") as f:
            return json.load(f)
    return json.load(sys.stdin)


if __name__ == "__main__":
    data = _load_input()
    role_name = data["role"]["name"]
    out_path = _default_output_path(role_name)
    saved = build_jd(data, out_path)
    print(f"WROTE: {saved}")
