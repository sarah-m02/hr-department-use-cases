"""
Locked template for the PIF-styled Candidate Assessment artifacts.

Renders TWO files from a single JSON blob:
  1. Assessment Rubric — Word doc (rubric + methodology)
  2. Scorecard — Excel (Ranked Scorecard + Per-Candidate Evidence)

Same code every run means pixel-identical formatting — fonts, colors,
column widths, table borders, tier colors, spacing all baked in.

Called by the hr-candidate-assessment skill (v1.5.0+) at Step 8. Do NOT
invoke the docx or xlsx skills for the render step and do NOT hand-generate
python-docx / openpyxl code. Assemble a JSON blob matching the schema below
and run this script.

JSON schema:

{
  "role": {
    "name": "Senior Analyst",
    "role_slug": "Senior_Analyst",
    "level": "Senior Analyst / Manager (mid-career, 4-8 yrs)",
    "division": "Corporate function — Risk Management",
    "summary_paragraph": "Role summary + scoring emphasis + must-have threshold, in one paragraph."
  },
  "dimensions": [
    {"name": "Functional & Technical Fit", "weight_pct": 30, "description": "..."},
    {"name": "Relevant Experience",        "weight_pct": 25, "description": "..."},
    {"name": "Leadership & Ownership",     "weight_pct": 20, "description": "..."},
    {"name": "Domain / Sector Alignment",  "weight_pct": 15, "description": "..."},
    {"name": "Culture & Values Fit",       "weight_pct": 10, "description": "..."}
  ],
  "must_haves": {
    "threshold_note": "Threshold: ...",
    "requirements": ["req 1", "req 2", ...]
  },
  "tiers": [
    {"name": "Advance to interview",  "range": "≥ 75, all must-haves met",   "recommendation": "..."},
    {"name": "Second-look pool",      "range": "55–74 (...)",                 "recommendation": "..."},
    {"name": "Not moving forward",    "range": "< 55, or must-have failure",  "recommendation": "..."}
  ],
  "blind_review_note": "Before scoring, each CV was stripped ...",
  "batch_summary": ["line 1", "line 2", "line 3"],
  "candidates": [
    {
      "rank": 1,
      "name": "Sarah Al-Mutairi",
      "tier": "Advance to interview",
      "composite": 85,
      "scores": [
        {"dim_name": "Functional & Technical Fit", "raw": 88, "weighted": 26.4, "evidence": "..."},
        ...
      ],
      "must_have_status": "Pass — all must-haves met",
      "red_flags": "None",
      "key_strengths": "...",
      "key_gaps": "...",
      "recommendation": "Advance — ..."
    },
    ...
  ]
}
"""
from __future__ import annotations

import json
import os
import sys
from datetime import datetime
from pathlib import Path
from typing import Tuple

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from openpyxl import Workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter


# ---------- Shared style tokens (matches build_jd.py) ----------

PIF_GREEN = "005C4D"
TAN = "C4984F"
TEXT_GRAY = "595959"
SOFT_GRAY = "9A9A9A"
LIGHT_GRAY = "F2F2F2"
WHITE = "FFFFFF"

FONT_PRIMARY = "Fund Light"

# Docx sizes
TITLE_PT = 20
SUBTITLE_PT = 12
SECTION_HEADING_PT = 14
BODY_PT = 11
TABLE_HEADER_PT = 11
TABLE_BODY_PT = 10
FOOTER_PT = 8


# ---------- File-lock retry ----------

def _save_with_retry(save_fn, base_path: Path) -> Path:
    """Try base_path, then _v2..._v9 suffixes if a PermissionError fires."""
    candidates = [base_path] + [
        base_path.with_name(f"{base_path.stem}_v{i}{base_path.suffix}")
        for i in range(2, 10)
    ]
    for p in candidates:
        try:
            save_fn(p)
            return p
        except PermissionError:
            continue
    raise PermissionError(f"All {len(candidates)} candidate paths for {base_path.name} are locked.")


# ---------- Docx helpers ----------

def _set_run_font(run, name: str = FONT_PRIMARY, size: int = BODY_PT,
                  color_hex: str = TEXT_GRAY, bold: bool = False, italic: bool = False) -> None:
    run.font.name = name
    r = run._element
    rPr = r.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:ascii"), name)
    rFonts.set(qn("w:hAnsi"), name)
    rFonts.set(qn("w:cs"), name)
    rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color_hex)
    run.bold = bold
    run.italic = italic


def _add_para(doc, text: str, size: int = BODY_PT, color: str = TEXT_GRAY,
              bold: bool = False, italic: bool = False, align=None, space_after: int = 6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run(text)
    _set_run_font(run, size=size, color_hex=color, bold=bold, italic=italic)
    return p


def _add_heading(doc, text: str, size: int = SECTION_HEADING_PT):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    _set_run_font(run, size=size, color_hex=PIF_GREEN, bold=True)
    return p


def _add_bullet(doc, text: str, size: int = BODY_PT, color: str = TEXT_GRAY):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.4)
    run = p.add_run(text)
    _set_run_font(run, size=size, color_hex=color)
    return p


def _shade_cell(cell, hex_color: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _set_cell_borders(cell, color: str = SOFT_GRAY, size_eighths: int = 4) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), str(size_eighths))
        b.set(qn("w:color"), color)
        tcBorders.append(b)
    tcPr.append(tcBorders)


def _horizontal_rule(doc, color: str = PIF_GREEN) -> None:
    p = doc.add_paragraph()
    pPr = p._element.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "18")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(8)


# ---------- Rubric docx builder ----------

def build_rubric_docx(data: dict, out_path: Path) -> Path:
    doc = Document()
    for section in doc.sections:
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)

    role = data["role"]

    # Header
    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run_t = title.add_run(f"Candidate Assessment Rubric — {role['name']}")
    _set_run_font(run_t, size=TITLE_PT, color_hex=PIF_GREEN, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(4)
    run_s = subtitle.add_run(f"{role['division']} · {role['level']} · Alat")
    _set_run_font(run_s, size=SUBTITLE_PT, color_hex=TEXT_GRAY, italic=True)

    _horizontal_rule(doc)

    # Role Summary
    _add_heading(doc, "Role Summary")
    _add_para(doc, role["summary_paragraph"], space_after=8)

    # Rubric Dimensions table
    _add_heading(doc, "Rubric Dimensions")
    dims = data["dimensions"]
    tbl = doc.add_table(rows=1 + len(dims), cols=3)
    col_widths = [Cm(6.5), Cm(2.5), Cm(8.8)]
    for col_idx, w in enumerate(col_widths):
        for row in tbl.rows:
            row.cells[col_idx].width = w

    for i, h in enumerate(["Dimension", "Weight", "Description"]):
        cell = tbl.rows[0].cells[i]
        _shade_cell(cell, PIF_GREEN)
        _set_cell_borders(cell, color=PIF_GREEN, size_eighths=6)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        _set_run_font(run, size=TABLE_HEADER_PT, color_hex=WHITE, bold=True)

    for r_idx, d in enumerate(dims, start=1):
        fill = LIGHT_GRAY if r_idx % 2 == 0 else WHITE
        values = [d["name"], f"{d['weight_pct']}%", d["description"]]
        for c_idx, v in enumerate(values):
            cell = tbl.rows[r_idx].cells[c_idx]
            _shade_cell(cell, fill)
            _set_cell_borders(cell, color=SOFT_GRAY, size_eighths=4)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(v)
            _set_run_font(run, size=TABLE_BODY_PT, color_hex=TEXT_GRAY,
                          bold=(c_idx == 0))

    # Must-Have Requirements
    _add_heading(doc, "Must-Have Requirements")
    _add_para(doc, data["must_haves"]["threshold_note"], italic=True, space_after=6)
    for req in data["must_haves"]["requirements"]:
        _add_bullet(doc, req)

    # Tier Definitions table
    _add_heading(doc, "Tier Definitions")
    tiers = data["tiers"]
    tbl2 = doc.add_table(rows=1 + len(tiers), cols=3)
    tier_col_widths = [Cm(5.0), Cm(6.0), Cm(6.8)]
    for col_idx, w in enumerate(tier_col_widths):
        for row in tbl2.rows:
            row.cells[col_idx].width = w

    for i, h in enumerate(["Tier", "Score Range", "Recommendation"]):
        cell = tbl2.rows[0].cells[i]
        _shade_cell(cell, PIF_GREEN)
        _set_cell_borders(cell, color=PIF_GREEN, size_eighths=6)
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(h)
        _set_run_font(run, size=TABLE_HEADER_PT, color_hex=WHITE, bold=True)

    tier_colors = {
        "Advance to interview": PIF_GREEN,
        "Second-look pool": TAN,
        "Not moving forward": SOFT_GRAY,
    }
    for r_idx, t in enumerate(tiers, start=1):
        for c_idx, v in enumerate([t["name"], t["range"], t["recommendation"]]):
            cell = tbl2.rows[r_idx].cells[c_idx]
            _set_cell_borders(cell, color=SOFT_GRAY, size_eighths=4)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(v)
            if c_idx == 0:
                # Tier label cell — colored fill, white bold
                _shade_cell(cell, tier_colors.get(t["name"], TEXT_GRAY))
                _set_run_font(run, size=TABLE_BODY_PT, color_hex=WHITE, bold=True)
            else:
                fill = LIGHT_GRAY if r_idx % 2 == 0 else WHITE
                _shade_cell(cell, fill)
                _set_run_font(run, size=TABLE_BODY_PT, color_hex=TEXT_GRAY)

    # Blind Review Protocol
    _add_heading(doc, "Blind Review Protocol")
    _add_para(doc, data["blind_review_note"], space_after=8)

    # Batch Summary
    _add_heading(doc, "Batch Summary")
    for line in data["batch_summary"]:
        _add_para(doc, line, space_after=4)

    # Footer
    section = doc.sections[0]
    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.text = ""
    run = fp.add_run("Confidential — HR use only")
    _set_run_font(run, size=FOOTER_PT, color_hex=SOFT_GRAY, italic=True)

    def _save(p: Path) -> None:
        doc.save(str(p))
    return _save_with_retry(_save, out_path)


# ---------- Scorecard xlsx builder ----------

def _xl_font(color_hex: str = TEXT_GRAY, bold: bool = False, size: int = 11) -> Font:
    return Font(name=FONT_PRIMARY, color=color_hex, bold=bold, size=size)


def _xl_fill(color_hex: str) -> PatternFill:
    return PatternFill(start_color=color_hex, end_color=color_hex, fill_type="solid")


def _xl_border() -> Border:
    thin = Side(border_style="thin", color=SOFT_GRAY)
    return Border(left=thin, right=thin, top=thin, bottom=thin)


TIER_FILLS = {
    "Advance to interview": PIF_GREEN,
    "Second-look pool": TAN,
    "Not moving forward": SOFT_GRAY,
}


def build_scorecard_xlsx(data: dict, out_path: Path) -> Path:
    wb = Workbook()
    _build_ranked_sheet(wb, data)
    _build_evidence_sheet(wb, data)

    def _save(p: Path) -> None:
        wb.save(str(p))
    return _save_with_retry(_save, out_path)


def _build_ranked_sheet(wb: Workbook, data: dict) -> None:
    ws = wb.active
    ws.title = "Ranked Scorecard"
    ws.freeze_panes = "A2"

    # Weight lookup: dim_name → weight_pct
    dim_weights = {d["name"]: d["weight_pct"] for d in data["dimensions"]}

    headers = [
        "Rank", "Candidate", "Tier", "Composite",
        f"Technical ({dim_weights.get('Functional & Technical Fit', 30)}%)",
        f"Experience ({dim_weights.get('Relevant Experience', 25)}%)",
        f"Leadership ({dim_weights.get('Leadership & Ownership', 20)}%)",
        f"Domain ({dim_weights.get('Domain / Sector Alignment', 15)}%)",
        f"Culture ({dim_weights.get('Culture & Values Fit', 10)}%)",
        "Must-Have Status", "Red Flags", "Key Strengths", "Key Gaps", "Recommendation",
    ]
    col_widths = [6, 22, 22, 12, 16, 16, 16, 14, 14, 30, 30, 40, 40, 40]

    for col_i, (h, w) in enumerate(zip(headers, col_widths), start=1):
        cell = ws.cell(row=1, column=col_i, value=h)
        cell.font = _xl_font(color_hex=WHITE, bold=True, size=11)
        cell.fill = _xl_fill(PIF_GREEN)
        cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        cell.border = _xl_border()
        ws.column_dimensions[get_column_letter(col_i)].width = w
    ws.row_dimensions[1].height = 24

    # Order the score dimensions in the same order as the header columns
    dim_order = [
        "Functional & Technical Fit",
        "Relevant Experience",
        "Leadership & Ownership",
        "Domain / Sector Alignment",
        "Culture & Values Fit",
    ]

    for row_i, c in enumerate(data["candidates"], start=2):
        fill_hex = LIGHT_GRAY if row_i % 2 == 0 else None
        scores_by_dim = {s["dim_name"]: s for s in c["scores"]}

        weighted_vals = [scores_by_dim.get(n, {}).get("weighted") for n in dim_order]

        row_values = [
            c["rank"],
            c["name"],
            c["tier"],
            c["composite"],
            *weighted_vals,
            c.get("must_have_status", ""),
            c.get("red_flags", ""),
            c.get("key_strengths", ""),
            c.get("key_gaps", ""),
            c.get("recommendation", ""),
        ]

        for col_i, v in enumerate(row_values, start=1):
            cell = ws.cell(row=row_i, column=col_i, value=v)
            cell.font = _xl_font(size=10)
            cell.border = _xl_border()
            cell.alignment = Alignment(vertical="top", wrap_text=True)
            if fill_hex is not None:
                cell.fill = _xl_fill(fill_hex)

        # Tier cell (column 3) — colored fill + white bold text
        tier_cell = ws.cell(row=row_i, column=3)
        tier_cell.fill = _xl_fill(TIER_FILLS.get(c["tier"], TEXT_GRAY))
        tier_cell.font = _xl_font(color_hex=WHITE, bold=True, size=10)
        tier_cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)

        # Composite score cell — bold
        comp_cell = ws.cell(row=row_i, column=4)
        comp_cell.font = _xl_font(size=11, bold=True)
        comp_cell.alignment = Alignment(horizontal="center", vertical="center")


def _build_evidence_sheet(wb: Workbook, data: dict) -> None:
    ws = wb.create_sheet("Per-Candidate Evidence")

    col_widths = [30, 15, 80]
    for col_i, w in enumerate(col_widths, start=1):
        ws.column_dimensions[get_column_letter(col_i)].width = w

    # Top banner
    ws.merge_cells("A1:C1")
    banner = ws.cell(row=1, column=1,
                     value="Per-Candidate Evidence — auditable score trail")
    banner.font = _xl_font(color_hex=WHITE, bold=True, size=12)
    banner.fill = _xl_fill(PIF_GREEN)
    banner.alignment = Alignment(horizontal="left", vertical="center")
    banner.border = _xl_border()
    ws.row_dimensions[1].height = 24

    row = 2
    for c in data["candidates"]:
        # Candidate header row (merged)
        ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=3)
        header = ws.cell(row=row, column=1,
                         value=f"{c['name']} — {c['tier']} — Composite {c['composite']}")
        header.font = _xl_font(color_hex=WHITE, bold=True, size=11)
        header.fill = _xl_fill(PIF_GREEN)
        header.alignment = Alignment(horizontal="left", vertical="center")
        header.border = _xl_border()
        ws.row_dimensions[row].height = 22
        row += 1

        # Sub-header
        for col_i, h in enumerate(["Dimension", "Score (raw)", "CV Evidence"], start=1):
            cell = ws.cell(row=row, column=col_i, value=h)
            cell.font = _xl_font(bold=True, size=10)
            cell.fill = _xl_fill(LIGHT_GRAY)
            cell.border = _xl_border()
            cell.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        row += 1

        # Dimension rows
        for i, score in enumerate(c["scores"]):
            fill_hex = None if i % 2 == 0 else LIGHT_GRAY
            values = [score["dim_name"], score["raw"], score.get("evidence", "")]
            for col_i, v in enumerate(values, start=1):
                cell = ws.cell(row=row, column=col_i, value=v)
                cell.font = _xl_font(size=10, bold=(col_i == 1))
                cell.border = _xl_border()
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                if fill_hex is not None:
                    cell.fill = _xl_fill(fill_hex)
                if col_i == 2:
                    cell.alignment = Alignment(horizontal="center", vertical="center")
            row += 1

        # Blank spacer row
        row += 1


# ---------- Public entry point ----------

def build_assessment(data: dict, out_dir: Path) -> Tuple[Path, Path]:
    out_dir.mkdir(parents=True, exist_ok=True)
    today = datetime.now().strftime("%Y%m%d")
    slug = data["role"].get("role_slug") or _slug(data["role"]["name"])
    rubric_path = out_dir / f"{today}_Assessment_Rubric_{slug}.docx"
    scorecard_path = out_dir / f"{today}_Scorecard_{slug}.xlsx"

    rubric_saved = build_rubric_docx(data, rubric_path)
    scorecard_saved = build_scorecard_xlsx(data, scorecard_path)
    return rubric_saved, scorecard_saved


def _slug(name: str) -> str:
    return "".join(ch if ch.isalnum() else "_" for ch in name).strip("_")


def _default_out_dir() -> Path:
    return Path(os.path.expanduser("~")) / "HR-Workspace" / "hr-candidate-assessment" / "outputs"


def _load_input() -> dict:
    # utf-8-sig strips a BOM if present (Windows PowerShell writes utf-8 with BOM by default)
    if len(sys.argv) >= 2:
        with open(sys.argv[1], "r", encoding="utf-8-sig") as f:
            return json.load(f)
    return json.load(sys.stdin)


if __name__ == "__main__":
    data = _load_input()
    rubric, scorecard = build_assessment(data, _default_out_dir())
    print(f"WROTE_RUBRIC: {rubric}")
    print(f"WROTE_SCORECARD: {scorecard}")
